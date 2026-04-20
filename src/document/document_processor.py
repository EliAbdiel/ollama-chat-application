import fitz
import base64
import asyncio
import chainlit as cl
from io import BytesIO
from pathlib import Path
from docx import Document
from ollama import AsyncClient, ChatResponse
from src.log.logger import setup_logger
from typing import Optional, Dict, Any, Set, AsyncIterator
from src.document.processor_config import ProcessingConfig
from src.utils.config import OLLAMA_BASE_URL, OLLAMA_API_KEY


class DocumentProcessor:
    """
    Enhanced document processor with better error handling,
    support for more file types, and advanced features.
    """
    def __init__(
            self, 
            ollama_api_key: str = OLLAMA_API_KEY, 
            config: Optional[ProcessingConfig] = None, 
            host: str = OLLAMA_BASE_URL
        ) -> None:
        self.config = config or ProcessingConfig()
        self.client = AsyncClient(host=host, headers={'Authorization': f'Bearer {ollama_api_key}'})
        self.logger = setup_logger('DOCS PROCESSOR')
        self.file_processor_map = {
            '.pdf': self._extract_text_from_pdf_bytes,
            '.docx': self._extract_text_from_docx_bytes,
            '.txt': self._extract_text_from_txt_bytes,
            '.jpg': self._extract_content_from_image_bytes,
            '.jpeg': self._extract_content_from_image_bytes,
            '.png': self._extract_content_from_image_bytes,
        }

    def _read_bytes(self, file: cl.File) -> bytes:
        """Read file content as bytes, handling both in-memory and disk storage"""
        content = getattr(file, 'content', None)
        if isinstance(content, (bytes, bytearray)):
            return bytes(content)
        file_path = getattr(file, 'path', None)
        if isinstance(file_path, str) and file_path:
            with open(file_path, 'rb') as f:
                return f.read()
        raise ValueError("File content is not available as bytes or valid path")

    async def _get_expected_mime_types(self, extension: str) -> Set[str]:
        """Get expected MIME types for a given file extension"""
        return self.config.allowed_mime_types.get(extension.lower(), set())

    async def _get_file_info(self, filename: str, file_bytes: bytes, file_mime: str) -> Dict[str, Any]:
        """Extract file information and validate"""
        file_extension = Path(filename).suffix.lower()
        file_size = len(file_bytes)

        return {
            'filename': filename,
            'extension': file_extension,
            'size': file_size,
            'mime_type': file_mime
        }
    
    async def _validate_file(self, filename: str, file_bytes: bytes, file_mime: str) -> Dict[str, Any]:
        """Enhanced file validation with content type checking"""
        file_info = await self._get_file_info(filename=filename, file_bytes=file_bytes, file_mime=file_mime)

        # Validate file extension
        if file_info['extension'] not in self.config.allowed_extensions:
            raise ValueError(f"Unsupported file extension: {file_info['extension']}")

        # Validate file content matches extension
        expected_mimes = await self._get_expected_mime_types(extension=file_info['extension'])
        if expected_mimes and file_info['mime_type'] not in expected_mimes:
            raise ValueError(
                f"File content doesn't match extension: {file_info['extension']}. "
                f"Expected: {expected_mimes}, Got: {file_info['mime_type']}"
            )

        # Check file size
        if file_info['size'] > self.config.max_file_size:
            raise ValueError(f"File size ({file_info['size']} bytes) exceeds limit ({self.config.max_file_size} bytes)")

        # Validate filename for security
        if '..' in filename or filename.startswith('/'):
            raise ValueError("Invalid filename: potential path traversal detected")

        return file_info

    async def _clean_and_summarize_text(self, text: str, filename: str, doc_type: str = "document") -> str:
        """
        Clean and summarize text using Ollama.
        Enhanced with better prompts and error handling.
        """
        if not text or not text.strip():
            return "No extractable text content found."

        try:
            # Enhanced prompt for better summarization
            system_instruction = f"""You are a summarization assistant specialized in long-form content.

            Input context
            - **Filename**: "{filename}"
            - **Document Category**: "{doc_type}"

            When summarizing, you must:
            - Preserve original meaning, intent, and logical flow
            - Extract only the most relevant information
            - Remove redundancy while keeping factual accuracy
            - Use clear, structured, concise language

            Constraints:
            - No opinions, assumptions, or invented facts
            - No altering context beyond compression
            - If ambiguous, summarize only what is certain

            Tone & Style:
            - Neutral, professional, short, clear, direct
            - Summary length: 10–30% of original
            - Always end with a 1-sentence summary (≤20 words)"""

            result: ChatResponse = await self.client.chat(
                model=self.config.ollama_model,
                messages=[
                    {'role': 'system', 'content': system_instruction},
                    {'role': 'user', 'content': text}
                ],
                options={
                    'temperature': self.config.temperature,
                    'num_ctx': self.config.num_ctx,
                }
            )

            self.logger.info(f"Successfully processed {doc_type} text with Ollama")
            return result.message.content

        except Exception as e:
            self.logger.error(f"Error in text cleaning/summarization: {str(e)}")
            # Return original text if summarization fails
            return f"Original content:\n\n{text[:self.config.text_extract_limit]}"

    async def _extract_text_from_pdf_bytes(self, filename: str, pdf_bytes: bytes, pdf_mime: str) -> str:
        """
        Extract text from PDF bytes using PyMuPDF with enhanced features.
        """
        try:

            doc = fitz.open(stream=pdf_bytes, filetype='pdf')
            text = ''
            #metadata = {}

            # Extract metadata
            #metadata = doc.metadata
            self.logger.info(f"Processing PDF with {len(doc)} pages")

            # Extract text from each page
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"

            doc.close()

            # Log extraction summary
            self.logger.info(f"Extracted {len(text)} characters from PDF")

            result = await self._clean_and_summarize_text(text=text[:self.config.text_extract_limit], filename=filename, doc_type=pdf_mime)

            return result

        except Exception as e:
            self.logger.error(f"Error processing PDF: {str(e)}")
            raise ValueError(f"Failed to process PDF: {str(e)}")

    async def _extract_text_from_docx_bytes(self, filename: str, docx_bytes: bytes, docx_mime: str) -> str:
        """
        Extract text from DOCX bytes with enhanced features.
        """
        try:

            doc = Document(BytesIO(docx_bytes))
            text = ''

            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells if cell.text)
                    if row_text:
                        text += f"Table row: {row_text}\n"

            self.logger.info(f"Extracted {len(text)} characters from DOCX")

            result = await self._clean_and_summarize_text(text=text[:self.config.text_extract_limit].strip(), filename=filename, doc_type=docx_mime)

            return result

        except Exception as e:
            self.logger.error(f"Error processing DOCX: {str(e)}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")

    async def _extract_text_from_txt_bytes(self, filename: str, txt_bytes: bytes, txt_mime: str) -> str:
        """
        Extract text from plain text files.
        """
        try:

            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'windows-1252', 'iso-8859-1']
            text = None

            for encoding in encodings:
                try:
                    text = txt_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                # Fallback: replace errors
                text = txt_bytes.decode('utf-8', errors='replace')

            self.logger.info(f"Extracted {len(text)} characters from TXT")

            result = await self._clean_and_summarize_text(text=text[:self.config.text_extract_limit], filename=filename, doc_type=txt_mime)

            return result

        except Exception as e:
            self.logger.error(f"Error processing TXT: {str(e)}")
            raise ValueError(f"Failed to process text file: {str(e)}")

    async def _extract_content_from_image_bytes(self, filename: str, image_bytes: bytes, image_mime: str) -> str:
        """
        Extract content from image bytes using Ollama vision model.
        """
        try:

            image_base64 = base64.b64encode(image_bytes).decode('utf-8')

            # Enhanced prompt for better extraction
            prompt = """
                You are an expert vision-analysis assistant.  
                Analyze the image carefully and provide a structured, detailed report containing:

                1. **Extracted Text (OCR):**  
                - Transcribe all visible text exactly as it appears.  
                - Preserve line breaks, formatting, and labels when helpful.

                2. **Visual Description:**  
                - Describe all key elements in the image (objects, layout, colors, UI elements, diagrams, charts, people, etc.).  
                - Explain spatial relationships (e.g., “A banner at the top”, “A table with two columns”, “Buttons at the bottom”).

                3. **Key Information & Data:**  
                - Identify important information the image conveys (numbers, labels, steps, headings, metrics, actions, warnings, etc.).  
                - If the image is a document, form, screenshot, diagram, or code snippet, summarize its core content.

                4. **Context & Purpose:**  
                - Explain what the image is likely used for (e.g., instructions, a form to fill out, a dashboard, a configuration screen, a diagram explaining X).  
                - Provide potential intent or user action implied by the image.

                5. **Concise Summary:**  
                - A short, easy-to-understand summary of the most important information from the image.

                **Guidelines:**  
                - Be objective: do not hallucinate nonexistent text or content.  
                - If a section has no information, state “None found”.  
                - Make the response organized using headings and bullet points.

                The image will be provided separately.
                """

            messages = [
                {
                    'role': 'user',
                    'content': prompt,
                    'images': [image_base64]
                }
            ]

            response: ChatResponse = await self.client.chat(
                model=self.config.vision_model,
                messages=messages,
                think=True,
                options={
                    'temperature': self.config.temperature,
                    'num_ctx': self.config.num_ctx,
                }
            )

            self.logger.info("Successfully processed image with vision model")

            result = await self._clean_and_summarize_text(text=response.message.content, filename=filename, doc_type=image_mime)

            return result

        except Exception as e:
            self.logger.error(f"Error processing image: {str(e)}")
            raise ValueError(f"Failed to process image: {str(e)}")
    
    async def process_document_async(self, filename: str, file_bytes: bytes, file_mime: str) -> str:
        """
        Process a document and extract structured content.

        Args:
            filename (str): The name of the file
            file_bytes (bytes): The raw bytes of the file
            file_mime (str): The mime type of the file

        Returns:
            str: Processed and summarized content
        """
        try:
            # Validate input
            if not filename or not file_bytes:
                raise ValueError("Filename and file_bytes are required")

            # Validate file and get info
            file_info = await self._validate_file(filename=filename, file_bytes=file_bytes, file_mime=file_mime)
            self.logger.info(f"Processing file: {file_info['filename']} "
                           f"({file_info['size']} bytes)")

            # Get appropriate processor
            processor = self.file_processor_map.get(file_info['extension'])
            if not processor:
                raise ValueError(f"Unsupported file extension: {file_info['extension']}")

            # Process the document
            result = await processor(filename, file_bytes, file_mime)

            self.logger.info(f"Successfully processed {filename}")
            return result

        except Exception as e:
            self.logger.error(f"Failed to process {filename}: {str(e)}")
            raise

    @cl.step(name="document", type="tool", show_input=False)
    async def process_multiple_files_async(self, files: list[cl.File]) -> dict:
        """
        Process multiple files and extract structured content.

        Args:
            files (list[cl.File]): List of files to process

        Returns:
            dict: Dictionary containing processed and summarized content
        """
        results = {
            "success": {},
            "failed": {}
        }

        if not files:
            return results

        async def process_single_file(file: cl.File):
            """
            Helper method for async processing of a single file
            """
            try:
                filename = str(file.name)
                file_bytes = self._read_bytes(file)
                file_mime = str(file.mime)

                self.logger.info(f"Processing file: {filename}, size: {len(file_bytes)} bytes, mime: {file_mime}")

                content = await self.process_document_async(filename, file_bytes, file_mime)

                results["success"][filename] = content

            except Exception as e:
                self.logger.error(f"Failed to process file: {str(e)}")
                results["failed"]["Status"] = str(e)

        tasks = []
        for file in files:
            current_task = asyncio.create_task(process_single_file(file))
            tasks.append(current_task)

        async for task in asyncio.as_completed(tasks):
            await task
            self.logger.info("Completed processing file")

        return results

    @cl.step(name="summarize", type="tool", show_input=False)
    async def summarize_text(self, content: str) -> Optional[str]:
        """
        Summarize text using the document processor.

        Args:
            content (str): The content to summarize

        Returns:
            str: Summarized content
        """
        if not content:
            self.logger.warning("Content is empty")
            return None

        self.logger.info(f"Summarizing text with length: {len(content)}")
        return await self._clean_and_summarize_text(text=content, filename="Document", doc_type="document")  